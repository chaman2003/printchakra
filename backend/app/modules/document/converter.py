"""
File Format Converter Module
Supports conversion between various formats: JPG, PNG, PDF, DOCX
Uses PIL (Pillow) as primary method for maximum compatibility
"""

import io
import os
import traceback
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Inches
from PIL import Image

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("⚠️  pdf2image not available - PDF page extraction disabled")


class FileConverter:
    """Handle file format conversions"""

    SUPPORTED_FORMATS = ["jpg", "jpeg", "png", "pdf", "docx"]

    @staticmethod
    def get_file_extension(filename: str) -> str:
        """Get lowercase file extension without dot"""
        return os.path.splitext(filename)[1].lower().lstrip(".")

    @staticmethod
    def is_supported_format(format: str) -> bool:
        """Check if format is supported"""
        return format.lower() in FileConverter.SUPPORTED_FORMATS

    @staticmethod
    def convert_image_to_pdf(input_path: str, output_path: str) -> Tuple[bool, str]:
        """
        Convert image (JPG/PNG) to PDF using PIL
        Robust method that works with all image formats
        """
        try:
            print(f"🔄 Converting image to PDF: {os.path.basename(input_path)}")

            # Validate input file exists
            if not os.path.exists(input_path):
                return False, f"Input file not found: {input_path}"

            # Get file size
            file_size_mb = os.path.getsize(input_path) / (1024 * 1024)
            print(f"   Input file size: {file_size_mb:.2f} MB")

            # Open image and convert to RGB (required for PDF)
            with Image.open(input_path) as img:
                print(f"   Image format: {img.format}, Mode: {img.mode}, Size: {img.size}")

                # Convert any image mode to RGB
                if img.mode in ("RGBA", "LA", "P", "1", "L"):
                    # Create white background for transparency
                    rgb_img = Image.new("RGB", img.size, (255, 255, 255))

                    if img.mode == "P":
                        # Handle palette mode
                        img = img.convert("RGBA")

                    if img.mode == "RGBA" or img.mode == "LA":
                        # Paste with alpha channel
                        rgb_img.paste(img, mask=img.split()[-1])
                    else:
                        # Direct paste
                        rgb_img.paste(img)

                    img_to_save = rgb_img
                else:
                    # Already in RGB or compatible mode
                    img_to_save = img.convert("RGB") if img.mode != "RGB" else img

                # Save as PDF
                print(f"   Saving as PDF...")
                img_to_save.save(output_path, "PDF", quality=95, optimize=True)

            # Validate output
            if not os.path.exists(output_path):
                return False, "Output PDF file was not created"

            output_size_mb = os.path.getsize(output_path) / (1024 * 1024)
            print(f"✅ Successfully converted to PDF: {output_size_mb:.2f} MB")
            return True, f"Converted to PDF ({output_size_mb:.2f} MB)"

        except Exception as e:
            error_msg = f"Image to PDF conversion failed: {str(e)}"
            print(f"❌ {error_msg}")
            traceback.print_exc()
            return False, error_msg

    @staticmethod
    def convert_image_to_docx(input_path: str, output_path: str) -> Tuple[bool, str]:
        """Convert image (JPG/PNG) to DOCX"""
        try:
            # Create Word document
            doc = Document()

            # Get image dimensions
            with Image.open(input_path) as img:
                width, height = img.size
                aspect_ratio = height / width

            # Add image to document (max 6 inches wide)
            max_width = 6.0
            doc_width = min(max_width, width / 96)  # 96 DPI
            doc.add_picture(input_path, width=Inches(doc_width))

            # Add filename as caption
            filename = os.path.basename(input_path)
            doc.add_paragraph(f"Image: {filename}", style="Caption")

            # Save document
            doc.save(output_path)

            print(f"✅ Converted image to DOCX: {output_path}")
            return True, f"Successfully converted to DOCX"

        except Exception as e:
            error_msg = f"Image to DOCX conversion failed: {str(e)}"
            print(f"❌ {error_msg}")
            traceback.print_exc()
            return False, error_msg

    @staticmethod
    def convert_image_format(
        input_path: str, output_path: str, target_format: str
    ) -> Tuple[bool, str]:
        """Convert between image formats (JPG ↔ PNG)"""
        try:
            with Image.open(input_path) as img:
                # Convert RGBA to RGB for JPEG
                if target_format.lower() in ["jpg", "jpeg"] and img.mode in ("RGBA", "LA", "P"):
                    rgb_img = Image.new("RGB", img.size, (255, 255, 255))
                    if img.mode == "P":
                        img = img.convert("RGBA")
                    rgb_img.paste(img, mask=img.split()[3] if img.mode == "RGBA" else None)
                    img = rgb_img

                # Save in target format
                if target_format.lower() in ["jpg", "jpeg"]:
                    img.save(output_path, "JPEG", quality=95, optimize=True)
                elif target_format.lower() == "png":
                    img.save(output_path, "PNG", optimize=True)
                else:
                    img.save(output_path, target_format.upper())

            print(f"✅ Converted image format: {output_path}")
            return True, f"Successfully converted to {target_format.upper()}"

        except Exception as e:
            error_msg = f"Image format conversion failed: {str(e)}"
            print(f"❌ {error_msg}")
            traceback.print_exc()
            return False, error_msg

    @staticmethod
    def convert_pdf_to_images(
        input_path: str, output_dir: str, format: str = "jpg"
    ) -> Tuple[bool, str, List[str]]:
        """
        Convert PDF to images (one image per page)
        Note: This feature requires additional dependencies (pdf2image)
        Currently returning error to indicate feature not available
        """
        error_msg = "PDF to image conversion is not currently supported. Please install pdf2image package if needed."
        print(f"ℹ️ {error_msg}")
        return False, error_msg, []

    @staticmethod
    def convert_file(
        input_path: str, output_path: str, source_format: str, target_format: str
    ) -> Tuple[bool, str]:
        """
        Main conversion dispatcher
        Returns: (success, message)
        """
        source_format = source_format.lower()
        target_format = target_format.lower()

        # Check if conversion is needed
        if source_format == target_format:
            return False, "Source and target formats are the same"

        # Check if formats are supported
        if not FileConverter.is_supported_format(source_format):
            return False, f"Unsupported source format: {source_format}"
        if not FileConverter.is_supported_format(target_format):
            return False, f"Unsupported target format: {target_format}"

        # Image to PDF
        if source_format in ["jpg", "jpeg", "png"] and target_format == "pdf":
            return FileConverter.convert_image_to_pdf(input_path, output_path)

        # Image to DOCX
        elif source_format in ["jpg", "jpeg", "png"] and target_format == "docx":
            return FileConverter.convert_image_to_docx(input_path, output_path)

        # Image format conversion (JPG ↔ PNG)
        elif source_format in ["jpg", "jpeg", "png"] and target_format in ["jpg", "jpeg", "png"]:
            return FileConverter.convert_image_format(input_path, output_path, target_format)

        # PDF to image
        elif source_format == "pdf" and target_format in ["jpg", "jpeg", "png"]:
            output_dir = os.path.dirname(output_path)
            success, message, files = FileConverter.convert_pdf_to_images(
                input_path, output_dir, target_format
            )
            if success and files:
                # For single page PDFs, rename the first file to match output_path
                if len(files) == 1:
                    first_file = os.path.join(output_dir, files[0])
                    if first_file != output_path:
                        os.rename(first_file, output_path)
            return success, message

        # Unsupported conversion
        else:
            return False, f"Conversion from {source_format} to {target_format} is not supported"

    @staticmethod
    def batch_convert(
        input_files: List[str], output_dir: str, target_format: str
    ) -> Tuple[int, int, List[dict]]:
        """
        Batch convert multiple files
        Returns: (success_count, fail_count, results)
        """
        results = []
        success_count = 0
        fail_count = 0

        print(f"\n🔄 Batch Converting {len(input_files)} files to {target_format.upper()}...")

        for idx, input_file in enumerate(input_files, 1):
            try:
                print(f"\n   [{idx}/{len(input_files)}] Processing: {os.path.basename(input_file)}")

                # Validate input file exists
                if not os.path.exists(input_file):
                    fail_count += 1
                    results.append(
                        {
                            "input": os.path.basename(input_file),
                            "output": None,
                            "success": False,
                            "message": f"Input file not found: {input_file}",
                        }
                    )
                    print(f"   ❌ Input file not found")
                    continue

                # Get source format
                source_format = FileConverter.get_file_extension(input_file)
                print(f"       Source format: {source_format}")

                # Generate output filename
                base_name = os.path.splitext(os.path.basename(input_file))[0]
                output_filename = f"{base_name}.{target_format}"
                output_path = os.path.join(output_dir, output_filename)
                print(f"       Target: {output_filename}")

                # Convert
                success, message = FileConverter.convert_file(
                    input_file, output_path, source_format, target_format
                )

                if success:
                    success_count += 1
                    file_size = os.path.getsize(output_path) / (1024 * 1024)
                    results.append(
                        {
                            "input": os.path.basename(input_file),
                            "output": output_filename,
                            "success": True,
                            "message": message,
                            "size_mb": round(file_size, 2),
                        }
                    )
                    print(f"       ✅ Success ({file_size:.2f} MB): {message}")

                    # Extract PDF pages if target format is PDF
                    if target_format.lower() == 'pdf':
                        extract_success, extracted_pages, extract_msg = FileConverter.extract_pdf_pages(
                            output_path, output_dir
                        )
                        if extract_success:
                            print(f"       📄 {extract_msg}")
                        else:
                            print(f"       ⚠️  Page extraction failed: {extract_msg}")

                else:
                    fail_count += 1
                    results.append(
                        {
                            "input": os.path.basename(input_file),
                            "output": None,
                            "success": False,
                            "message": message,
                        }
                    )
                    print(f"       ❌ Failed: {message}")

            except Exception as e:
                fail_count += 1
                error_msg = str(e)
                results.append(
                    {
                        "input": os.path.basename(input_file),
                        "output": None,
                        "success": False,
                        "message": error_msg,
                    }
                )
                print(f"       ❌ Exception: {error_msg}")
                traceback.print_exc()

        print(f"\n📊 Batch conversion complete: {success_count} succeeded, {fail_count} failed\n")
        return success_count, fail_count, results

    @staticmethod
    def merge_images_to_pdf(input_files: List[str], output_path: str) -> Tuple[bool, str]:
        """
        Merge multiple images into a single PDF using PIL
        Robust method that handles all image types
        """
        try:
            if not input_files:
                return False, "No input files provided"

            print(f"\n🔄 Merging {len(input_files)} images into single PDF...")

            # Validate all input files exist
            for i, f in enumerate(input_files, 1):
                if not os.path.exists(f):
                    return False, f"Input file #{i} not found: {f}"
                print(f"   [{i}/{len(input_files)}] ✓ {os.path.basename(f)}")

            # Process and convert all images to RGB
            image_list = []

            for idx, input_file in enumerate(input_files, 1):
                try:
                    print(
                        f"   Processing [{idx}/{len(input_files)}]: {os.path.basename(input_file)}"
                    )

                    with Image.open(input_file) as img:
                        print(f"      Format: {img.format}, Mode: {img.mode}, Size: {img.size}")

                        # Convert any mode to RGB
                        if img.mode in ("RGBA", "LA", "P", "1", "L"):
                            # Create white background
                            rgb_img = Image.new("RGB", img.size, (255, 255, 255))

                            if img.mode == "P":
                                img = img.convert("RGBA")

                            if img.mode in ("RGBA", "LA"):
                                rgb_img.paste(img, mask=img.split()[-1])
                            else:
                                rgb_img.paste(img)

                            image_list.append(rgb_img)
                        else:
                            image_list.append(img.convert("RGB"))

                except Exception as e:
                    print(f"      ⚠️ Error processing image: {e}")
                    return False, f"Failed to process {os.path.basename(input_file)}: {str(e)}"

            if not image_list:
                return False, "No valid images to merge"

            print(f"   Saving {len(image_list)} images to PDF...")

            # Save all images as a single PDF
            # First image is the main image, rest are appended
            first_image = image_list[0]
            remaining_images = image_list[1:] if len(image_list) > 1 else []

            first_image.save(
                output_path,
                "PDF",
                save_all=True,
                append_images=remaining_images,
                quality=95,
                optimize=True,
            )

            # Validate output
            if not os.path.exists(output_path):
                return False, "Output PDF file was not created"

            output_size_mb = os.path.getsize(output_path) / (1024 * 1024)
            print(
                f"✅ Successfully merged {len(image_list)} images into PDF ({output_size_mb:.2f} MB)"
            )
            return True, f"Merged {len(image_list)} images into PDF ({output_size_mb:.2f} MB)"

        except Exception as e:
            error_msg = f"PDF merge failed: {str(e)}"
            print(f"❌ {error_msg}")
            traceback.print_exc()
            return False, error_msg

    @staticmethod
    def extract_pdf_pages(pdf_path: str, output_dir: str) -> Tuple[bool, List[str], str]:
        """
        Extract all pages from a PDF as images
        Returns: (success, list_of_image_paths, message)
        """
        if not PDF2IMAGE_AVAILABLE:
            return False, [], "pdf2image library not available"

        try:
            if not os.path.exists(pdf_path):
                return False, [], f"PDF file not found: {pdf_path}"

            print(f"\n📄 Extracting pages from PDF: {os.path.basename(pdf_path)}")

            # Create pages subdirectory
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            pages_dir = os.path.join(output_dir, f"{base_name}_pages")
            os.makedirs(pages_dir, exist_ok=True)

            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=150)  # 150 DPI for good quality/size tradeoff

            print(f"   Total pages: {len(images)}")

            extracted_pages = []
            blank_count = 0

            for page_num, image in enumerate(images, 1):
                # Check if page is blank (>95% white)
                pixels = list(image.convert('RGB').getdata())
                total_pixels = len(pixels)
                light_pixels = sum(1 for r, g, b in pixels if r > 200 and g > 200 and b > 200)
                light_ratio = light_pixels / total_pixels if total_pixels > 0 else 0

                if light_ratio > 0.95:
                    print(f"   ⊘ Page {page_num}: Skipped (blank)")
                    blank_count += 1
                    continue

                # Save page as JPEG
                page_filename = f"{base_name}_page_{page_num:03d}.jpg"
                page_path = os.path.join(pages_dir, page_filename)
                image.save(page_path, 'JPEG', quality=85, optimize=True)

                extracted_pages.append(page_path)
                file_size = os.path.getsize(page_path) / 1024  # KB
                print(f"   ✓ Page {page_num}: Extracted ({file_size:.1f} KB)")

            if not extracted_pages:
                return False, [], f"No valid pages extracted (all {blank_count} pages were blank)"

            print(f"   ✅ Extracted {len(extracted_pages)} pages (skipped {blank_count} blank pages)")
            return True, extracted_pages, f"Extracted {len(extracted_pages)} pages"

        except Exception as e:
            error_msg = f"PDF page extraction failed: {str(e)}"
            print(f"   ❌ {error_msg}")
            traceback.print_exc()
            return False, [], error_msg
